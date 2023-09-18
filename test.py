# Import docx NOT python-docx
from tkinter import *
from tkinter.ttk import Frame, Button, Label
import docx
from docx.shared import Inches
import os
from os import listdir
from PIL import Image
import datetime

def createDoc(inpTitle, inpFile):
    list = []

    # get the path or directory
    folder_dir = os.getcwd()
    directory = 'Pics'
    path = os.path.join(folder_dir, directory)
    if os.path.exists(path) == False:
        os.makedirs(path)
    
    for images in os.listdir(folder_dir):
        # check if the image ends with png or jpg or jpeg
        if (images.endswith(".png") or images.endswith(".jpg") or images.endswith(".jpeg")):
            # display
            img = Image.open(images)
            img.save(os.path.join(path, images))
            print(images)
            list.append(images)
            #img.close()
    
    # Create an instance of a word document
    doc = docx.Document()
    
    # Add a Title to the document
    if inpTitle == "": 
        doc.add_heading('Report', 0)
    else:
        doc.add_heading(inpTitle, 0) 
    
    index=1
    for i in list:
        img = Image.open(i)
        if img.size[0] > 500: 
            doc.add_picture(i, width=Inches(5.0))
        elif img.size[1] > 300:
            doc.add_picture(i, height=Inches(3.5))
        else: doc.add_picture(i)
        img.close()
        os.remove(i)
        
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Light List Accent 6'
        cell_a = table.cell(0,0)
        cell_b = table.cell(1,0)
        picTxt = cell_a.text = 'Picture ' + str(index) 
        cell_b.style = 'Normal'
        cell_b.text = 'Info: '
        index+=1
        doc.add_paragraph()
    
    # Now save the document to a location
    # creating a unique name for the file based on date and time
    now = datetime.datetime.now()
    now = str(now.year) + str(now.month) + str(now.day) + str(now.hour) + str(now.minute) + str(now.second) + "."
    if inpFile != "":
        file_name = str(now) + inpFile + ".report." + "docx"
    else: file_name = str(now) + "report." + "docx"
    print(file_name)
    print(now)
    doc.save(file_name)



def createWindow():
   
    def fileName():
        l1.pack_forget()
        btn1.forget()
        btn2.pack_forget()
        btn3.pack()
        yes.pack()
        print(inpTitle)
        
    def saved():
        #l1_1.pack_forget()
        l2.pack_forget()
        l5.pack_forget()
        btn3.pack_forget()
        save.pack()
        inputTitle.pack_forget()
        inputFile.pack_forget()
        inpTitle = inputTitle.get(1.0, "end-1c")
        inpFile = inputFile.get(1.0, "end-1c")
        createDoc(inpTitle, inpFile)
        
    
    def takeDefault():
        btn1.forget()
        btn2.pack_forget()
        l1.pack_forget()
        no.pack()
        createDoc('')
        
    # Creating tkinter window with fixed geometry
    root = Tk()
    root.geometry('400x400')
    inpTitle = ""
    inpFile = ""

    #main = Frame(root)
    yes = Frame(root)
    save = Frame(root)
    no = Frame(root)
    
    msg1_1 ='You made it  so far! Go you!'
    
    l1_1 = Label(yes, text=msg1_1)
    l1_1.place(x=80, y=30)
    l1_1.pack()
    
    # This will create a LabelFrame
    l1 = Label(root, text='Would you like a default name for the document')
    l1.place(x=80, y=30)
    l1.pack()
    
    msg3 = \
    '''
    
          Your  file was created
    thank you  for levereging this tool 
          to enhance your skills
    '''
    
    l3 = Label(save, text=msg3)
    l3.place(x=30, y=30)
    l3.pack()
    
    # Buttons
    btn3 = Button(yes, text='save', command=saved)
    btn3.place(x=30, y=30)
    btn3.pack()
    
    btn1 = Button(root, text='yes', command=fileName)
    btn1.place(x=30, y=30)
    btn1.pack()
    
    btn2 = Button(root, text='no', command=takeDefault)
    btn2.place(x=130, y=30)
    btn2.pack()
    
    msg2 = \
    '''
    Type the name you would like as a Title
    Then press save to enter! Doc Title not
    file name.
    '''
    
    l2 = Label(yes, text=msg2)
    l2.pack()
    
    inputTitle = Text(yes, height=1, width=30)
    inputTitle.pack()
    
    l5 = Label(yes, text="Enter descriptor for file name. Alpha only, please keep it simple")
    l5.pack()
    
    inputFile = Text(yes, height=1, width=30)
    inputFile.pack()
    #inputFile = str(input)
    
    # while inputFile.isalpha() != True:
    #     l5 = Label(yes, text="Try again you didnt enter ALPHA")
    #     l5.pack()
    #     inputFile = Text(yes, height=1, width=30)
    #     inputFile.pack() 
    #     inputFile = str(input)
    
    l4 = Label(no, text=msg3)
    l4.place(x=30, y=30)
    l4.pack()
    
    # This creates infinite loop which generally
    # waits for any interrupt (like keyboard or
    # mouse) to terminate
    root.mainloop()

createWindow()